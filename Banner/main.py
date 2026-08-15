"""LED 电子横幅主程序：模拟街头常见黑色点阵电子横幅。"""

import argparse
import json
import sys
import time
import traceback
import tkinter as tk
from pathlib import Path
from tkinter import filedialog, messagebox, ttk

from led_ticker import LEVEL_COLORS, render_preview_image, render_text_columns

DEFAULT_CONTENT = "你好，世界！"
CONFIG_NAME = "banner_config.json"
LOG_NAME = "banner.log"
MAX_CONTENT_CHARS = 3000


def parse_args(argv=None):
    p = argparse.ArgumentParser(description="LED 点阵电子横幅模拟器")
    p.add_argument("--content", default=None, help="横幅滚动内容（默认读配置，否则：你好，世界！）")
    p.add_argument("--rows", type=int, default=32, help="点阵行数（默认 32）")
    p.add_argument("--cols", type=int, default=192, help="点阵列数（默认 192）")
    p.add_argument("--cell", type=int, default=8, help="初始 LED 间距像素（默认 8）")
    p.add_argument("--fps", type=int, default=25, help="刷新率，帧/秒（默认 25）")
    p.add_argument("--speed", type=float, default=None, help="滚动速度，列/秒")
    p.add_argument("--direction", choices=("left", "right"), default=None, help="滚动方向")
    p.add_argument(
        "--preview",
        nargs="?",
        const="preview.png",
        metavar="PNG",
        help="不启动窗口，渲染一帧到 PNG 后退出",
    )
    p.add_argument(
        "--selftest",
        metavar="FILE",
        help="读取指定 Word/文本文件测试后退出（用于排查问题）",
    )
    return p.parse_args(argv)


def config_path():
    base = Path(sys.executable).parent if getattr(sys, "frozen", False) else Path(__file__).parent
    return base / CONFIG_NAME


def load_config():
    try:
        with open(config_path(), "r", encoding="utf-8") as f:
            return json.load(f)
    except (OSError, ValueError):
        return {}


def save_config(cfg):
    try:
        with open(config_path(), "w", encoding="utf-8") as f:
            json.dump(cfg, f, ensure_ascii=False, indent=2)
    except OSError:
        pass


def log_error(exc):
    """把异常写入 exe/脚本所在目录的 banner.log，便于排查。"""
    try:
        base = Path(sys.executable).parent if getattr(sys, "frozen", False) else Path(__file__).parent
        with open(base / LOG_NAME, "a", encoding="utf-8") as f:
            f.write(f"\n[{time.strftime('%Y-%m-%d %H:%M:%S')}] {exc}\n")
            traceback.print_exc(file=f)
    except OSError:
        pass


def normalize_text(text):
    """把换行、连续空白折叠成单个空格（单行横幅）。"""
    return " ".join(str(text).split())


def prepare_text(text):
    """规范化并限制最大长度，避免超长文档导致点阵列数爆炸。"""
    return normalize_text(text)[:MAX_CONTENT_CHARS]


def extract_word_text(path):
    """读取 Word/文本文件内容，返回纯文本。"""
    p = Path(path)
    ext = p.suffix.lower()
    if ext == ".docx":
        try:
            import docx
        except ImportError:
            raise ValueError("缺少 python-docx 组件，无法读取 .docx 文档")
        doc = docx.Document(str(p))
        parts = [para.text for para in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        return "\n".join(parts)
    if ext == ".txt":
        return p.read_text(encoding="utf-8", errors="ignore")
    raise ValueError("仅支持 .docx（Word 文档）或 .txt 文本文件；.doc 请先用 Word 另存为 .docx")


def _to_hex(rgb):
    return "#%02x%02x%02x" % rgb


def enable_dpi_awareness():
    """启用 DPI 感知，避免 Windows 缩放导致界面文字模糊/马赛克。"""
    if sys.platform != "win32":
        return
    try:
        import ctypes

        try:
            ctypes.windll.shcore.SetProcessDpiAwareness(1)
        except Exception:
            ctypes.windll.user32.SetProcessDPIAware()
    except Exception:
        pass


def adjust_tk_scaling(root):
    """按真实 DPI 调整 Tk 文本缩放，保证文字清晰。"""
    try:
        scale = root.winfo_fpixels("1i") / 72.0
        if abs(scale - 1.0) > 0.05:
            root.tk.call("tk", "scaling", scale)
    except tk.TclError:
        pass


class LedBannerApp:
    PADDING = 18

    def __init__(self, root, text, rows, cols, cell, fps, speed, direction, on_change=None):
        self.root = root
        self.rows = rows
        self.cols = cols
        self.cell = max(int(cell), 4)
        self.interval = max(int(1000 / fps), 10)
        self.speed = max(5.0, float(speed))
        self.direction = direction if direction in ("left", "right") else "left"
        self.paused = False
        self.default_content = normalize_text(text) or DEFAULT_CONTENT
        self.word_path = None
        self.on_change = on_change
        self.settings_win = None
        self.canvas = None
        self._resize_job = None
        self.fill_colors = [_to_hex(c) for c in LEVEL_COLORS]

        self.columns = render_text_columns(self.default_content, rows=rows)
        self.text_cols = len(self.columns)
        self._reset_offset()
        self._force_redraw = True
        self.last_time = time.monotonic()

        root.title(self._title())
        root.configure(bg="black")
        base_w = cols * self.cell + self.PADDING * 2
        base_h = rows * self.cell + self.PADDING * 2
        root.minsize(cols * 4 + self.PADDING * 2, rows * 4 + self.PADDING * 2)
        root.geometry(f"{base_w}x{base_h}")
        root.update_idletasks()
        self._rebuild_grid()

        root.bind("<Configure>", self._on_configure)
        root.bind("<Button-1>", self.toggle_pause)
        root.bind("<Button-3>", self.open_settings)
        root.bind("<Escape>", lambda e: root.destroy())
        root.bind("<space>", self.toggle_pause)
        root.bind("+", self.speed_up)
        root.bind("=", self.speed_up)
        root.bind("-", self.speed_down)
        root.after(self.interval, self.tick)

    # ---------- 状态 ----------

    def _title(self):
        direction = "向左" if self.direction == "left" else "向右"
        title = f"LED 电子横幅 | {self.default_content} | 速度 {self.speed:.0f} | {direction}"
        if self.paused:
            title += "（已暂停）"
        return title

    def _reset_offset(self):
        if self.direction == "left":
            # 文字从屏幕右侧进入
            self.offset = float((self.text_cols - self.cols % self.text_cols) % self.text_cols)
        else:
            # 文字从屏幕左侧进入
            self.offset = 1.0

    def set_content(self, text):
        text = prepare_text(text) or DEFAULT_CONTENT
        self.columns = render_text_columns(text, rows=self.rows)
        self.text_cols = len(self.columns)
        self._reset_offset()
        self.prev_levels = [[0] * self.cols for _ in range(self.rows)]
        self._force_redraw = True
        self.root.title(self._title())

    def set_speed(self, value):
        self.speed = max(5.0, min(float(value), 200.0))
        self.root.title(self._title())

    def set_direction(self, direction):
        if direction in ("left", "right"):
            self.direction = direction
            self.root.title(self._title())

    # ---------- 窗口缩放（等比例放大/缩小，点阵铺满窗口） ----------

    def _on_configure(self, _event):
        if self._resize_job is not None:
            self.root.after_cancel(self._resize_job)
        self._resize_job = self.root.after(150, self._rebuild_grid)

    def _rebuild_grid(self):
        self._resize_job = None
        w = self.root.winfo_width()
        h = self.root.winfo_height()
        avail_w = w - self.PADDING * 2
        avail_h = h - self.PADDING * 2
        if avail_w < 40 or avail_h < 40:
            return
        cell = max(4, min(int(avail_w / self.cols), int(avail_h / self.rows), 32))
        if self.canvas is not None and cell == self.cell:
            return
        self.cell = cell
        grid_w = self.cols * cell
        grid_h = self.rows * cell
        ox = (w - grid_w) // 2
        oy = (h - grid_h) // 2

        if self.canvas is None:
            self.canvas = tk.Canvas(self.root, bg="black", highlightthickness=0)
            self.canvas.pack(fill="both", expand=True)
        else:
            self.canvas.delete("all")

        self.canvas.create_rectangle(
            ox - 2, oy - 2, ox + grid_w + 2, oy + grid_h + 2,
            outline="#3a3a3a", width=2,
        )
        led = max(cell - 2, 3)
        self.leds = []
        for r in range(self.rows):
            row_ids = []
            y = oy + r * cell
            for c in range(self.cols):
                x = ox + c * cell
                item = self.canvas.create_oval(
                    x, y, x + led, y + led,
                    fill=self.fill_colors[0], outline="",
                )
                row_ids.append(item)
            self.leds.append(row_ids)
        self.prev_levels = [[0] * self.cols for _ in range(self.rows)]
        self._force_redraw = True

    # ---------- 事件 ----------

    def toggle_pause(self, _event=None):
        self.paused = not self.paused
        self.root.title(self._title())

    def speed_up(self, _event=None):
        self.set_speed(self.speed + 5)

    def speed_down(self, _event=None):
        self.set_speed(self.speed - 5)

    def open_settings(self, _event=None):
        win = self.settings_win
        if win is not None:
            try:
                if win.win.winfo_exists():
                    win.win.lift()
                    win.win.focus_force()
                    return
            except tk.TclError:
                self.settings_win = None
        self.settings_win = SettingsDialog(self)

    # ---------- 渲染 ----------

    def tick(self):
        now = time.monotonic()
        dt = now - self.last_time
        self.last_time = now
        if not self.paused:
            mult = 1.0 if self.direction == "left" else -1.0
            self.offset += self.speed * dt * mult

        start = int(self.offset) % self.text_cols
        if self._force_redraw:
            # 内容切换/窗口重建后：无条件写入全部 LED，保证画面一致
            self._redraw_all(start)
        else:
            # 常规滚动：全量比对，只更新发生变化的 LED（每帧仅几十颗）
            self._redraw_diff(start)
        self._force_redraw = False
        self.root.after(self.interval, self.tick)

    def _redraw_all(self, start):
        for c in range(self.cols):
            col = self.columns[(start + c) % self.text_cols]
            for r in range(self.rows):
                v = col[r]
                self.canvas.itemconfig(self.leds[r][c], fill=self.fill_colors[v])
                self.prev_levels[r][c] = v

    def _redraw_diff(self, start):
        for c in range(self.cols):
            col = self.columns[(start + c) % self.text_cols]
            for r in range(self.rows):
                v = col[r]
                if v != self.prev_levels[r][c]:
                    self.canvas.itemconfig(self.leds[r][c], fill=self.fill_colors[v])
                    self.prev_levels[r][c] = v


class SettingsDialog:
    BG = "#1e1e1e"

    def __init__(self, app):
        self.app = app
        self.win = tk.Toplevel(app.root)
        self.win.title("横幅设置")
        self.win.configure(bg=self.BG)
        self.win.resizable(False, False)
        self.win.transient(app.root)
        self.win.protocol("WM_DELETE_WINDOW", self.close)

        style = ttk.Style(self.win)
        style.theme_use("clam")
        style.configure(".", background=self.BG, foreground="#e8e8e8", font=("Microsoft YaHei UI", 10))
        style.configure("TFrame", background=self.BG)
        style.configure("TLabel", background=self.BG, foreground="#e8e8e8")
        style.configure("Hint.TLabel", foreground="#9cc9f0")
        style.configure(
            "TLabelframe",
            background=self.BG,
            bordercolor="#3d3d3d",
            lightcolor="#3d3d3d",
            darkcolor="#3d3d3d",
        )
        style.configure("TLabelframe.Label", background=self.BG, foreground="#9cc9f0")
        style.configure("TButton", background="#3a3a3a", foreground="#fff", padding=(12, 6))
        style.map("TButton", background=[("active", "#4a4a4a"), ("pressed", "#2f2f2f")])
        style.configure("Accent.TButton", background="#2e6b34")
        style.map("Accent.TButton", background=[("active", "#3a8a42"), ("pressed", "#25592a")])
        style.configure("TRadiobutton", background=self.BG, foreground="#e8e8e8")
        style.map(
            "TRadiobutton",
            background=[("active", self.BG)],
            foreground=[("active", "#ffffff")],
        )
        style.configure(
            "TEntry",
            fieldbackground="#2a2a2a",
            foreground="#ffffff",
            insertcolor="#ffffff",
            bordercolor="#3d3d3d",
            lightcolor="#3d3d3d",
        )
        style.configure("TScale", background=self.BG, troughcolor="#3a3a3a")

        self.speed_var = tk.IntVar(value=int(app.speed))
        self.speed_text = tk.StringVar(value=f"{app.speed:.0f} 列/秒")
        self.dir_var = tk.StringVar(value=app.direction)
        self.content_var = tk.StringVar(value=app.default_content)
        word = app.word_path.name if app.word_path else "未选择 Word 文档"
        self.word_var = tk.StringVar(value=f"正在播放：{word}" if app.word_path else word)
        self.speed_var.trace_add(
            "write", lambda *_: self.speed_text.set(f"{self.speed_var.get()} 列/秒")
        )

        outer = ttk.Frame(self.win, padding=14)
        outer.pack(fill="both", expand=True)

        speed_box = ttk.LabelFrame(outer, text="滚动设置", padding=10)
        speed_box.pack(fill="x")
        row1 = ttk.Frame(speed_box)
        row1.pack(fill="x")
        ttk.Label(row1, text="移动速度").pack(side="left")
        ttk.Label(row1, textvariable=self.speed_text).pack(side="right")
        ttk.Scale(speed_box, from_=5, to=200, variable=self.speed_var).pack(
            fill="x", pady=(6, 2)
        )
        dir_row = ttk.Frame(speed_box)
        dir_row.pack(fill="x", pady=(8, 0))
        ttk.Label(dir_row, text="滚动方向").pack(side="left", padx=(0, 12))
        ttk.Radiobutton(dir_row, text="向左", value="left", variable=self.dir_var).pack(
            side="left"
        )
        ttk.Radiobutton(
            dir_row, text="向右", value="right", variable=self.dir_var
        ).pack(side="left", padx=(12, 0))

        content_box = ttk.LabelFrame(outer, text="内容设置", padding=10)
        content_box.pack(fill="x", pady=(12, 0))
        ttk.Label(content_box, text="默认显示内容").pack(anchor="w")
        self.content_entry = ttk.Entry(content_box, textvariable=self.content_var)
        self.content_entry.pack(fill="x", pady=(4, 8))
        btn_row = ttk.Frame(content_box)
        btn_row.pack(fill="x")
        ttk.Button(btn_row, text="选择 Word 文件播放…", command=self.pick_word).pack(
            side="left"
        )
        ttk.Button(btn_row, text="恢复默认内容", command=self.restore_default).pack(
            side="left", padx=(8, 0)
        )
        ttk.Label(
            content_box,
            textvariable=self.word_var,
            style="Hint.TLabel",
            wraplength=360,
            justify="left",
        ).pack(anchor="w", pady=(8, 0))

        bottom = ttk.Frame(outer)
        bottom.pack(fill="x", pady=(14, 0))
        ttk.Button(bottom, text="保存", style="Accent.TButton", command=self.save).pack(
            side="right"
        )
        ttk.Button(bottom, text="关闭", command=self.close).pack(side="right", padx=(0, 10))

    def close(self):
        self.app.settings_win = None
        self.win.destroy()

    def pick_word(self):
        path = filedialog.askopenfilename(
            parent=self.win,
            title="选择要播放的文档",
            filetypes=[
                ("Word 文档", "*.docx"),
                ("文本文件", "*.txt"),
                ("所有文件", "*.*"),
            ],
        )
        if not path:
            return
        try:
            text = extract_word_text(path)
        except Exception as exc:
            log_error(exc)
            messagebox.showerror(
                "无法读取文档",
                f"{exc}\n\n详细错误已写入 banner.log",
                parent=self.win,
            )
            return
        text = normalize_text(text)
        if not text:
            messagebox.showwarning(
                "文档为空",
                "该文档中没有可显示的文本（图片、文本框等内容无法读取）。",
                parent=self.win,
            )
            return
        truncated = len(text) > MAX_CONTENT_CHARS
        text = text[:MAX_CONTENT_CHARS]
        try:
            self.app.set_content(text)
        except Exception as exc:
            log_error(exc)
            messagebox.showerror("无法播送文档内容", str(exc), parent=self.win)
            return
        self.app.word_path = Path(path)
        label = f"正在播放：{self.app.word_path.name}"
        if truncated:
            label += "（内容过长，已截取前 3000 字）"
        self.word_var.set(label)

    def restore_default(self):
        content = self.content_var.get().strip() or DEFAULT_CONTENT
        self.app.word_path = None
        self.word_var.set("未选择 Word 文档")
        self.app.set_content(content)

    def save(self):
        speed = max(5, min(self.speed_var.get(), 200))
        direction = self.dir_var.get()
        content = self.content_var.get().strip() or DEFAULT_CONTENT
        self.app.set_speed(speed)
        self.app.set_direction(direction)
        self.app.default_content = content
        if self.app.word_path is None:
            self.app.set_content(content)
        if self.app.on_change:
            self.app.on_change({"content": content, "speed": speed, "direction": direction})
        self.close()


def main(argv=None):
    enable_dpi_awareness()
    args = parse_args(argv)
    if args.selftest:
        try:
            text = extract_word_text(args.selftest)
        except Exception as exc:
            print(f"SELFTEST_FAIL: {exc}")
            return 1
        norm = normalize_text(text)
        print(f"SELFTEST_OK chars={len(norm)} preview={norm[:80]}")
        return 0

    if args.preview:
        render_preview_image(
            prepare_text(args.content) or DEFAULT_CONTENT,
            rows=args.rows,
            cols=args.cols,
            out_path=args.preview,
        )
        print(f"预览已生成: {args.preview}")
        return 0

    cfg = load_config()
    content = args.content if args.content is not None else cfg.get("content", DEFAULT_CONTENT)
    speed = args.speed if args.speed is not None else float(cfg.get("speed", 30))
    direction = args.direction if args.direction is not None else cfg.get("direction", "left")

    root = tk.Tk()
    adjust_tk_scaling(root)
    LedBannerApp(
        root,
        content,
        args.rows,
        args.cols,
        args.cell,
        args.fps,
        speed,
        direction,
        on_change=save_config,
    )
    root.mainloop()
    return 0


if __name__ == "__main__":
    sys.exit(main())
